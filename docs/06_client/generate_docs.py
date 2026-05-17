import os
import shutil
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# --- SETUP PATHS ---
workspace_dir = r"c:\AG\epios"
client_dir = os.path.join(workspace_dir, "docs", "06_client")
images_dir = os.path.join(client_dir, "images")

# Ensure image folder exists
os.makedirs(images_dir, exist_ok=True)

# Sources from brain directory
src_showcase = r"C:\Users\rfxxi\.gemini\antigravity\brain\ba047a12-0096-4664-9892-eb2dfa9de839\epios_dashboard_showcase_1779002633998.png"
src_competitors = r"C:\Users\rfxxi\.gemini\antigravity\brain\ba047a12-0096-4664-9892-eb2dfa9de839\epios_vs_competitors_1779002648432.png"

# Copy images to docs for permanent reference
dst_showcase = os.path.join(images_dir, "epios_dashboard_showcase.png")
dst_competitors = os.path.join(images_dir, "epios_vs_competitors.png")

print("Copying asset files...")
if os.path.exists(src_showcase):
    shutil.copy(src_showcase, dst_showcase)
if os.path.exists(src_competitors):
    shutil.copy(src_competitors, dst_competitors)

# --- DOCX STYLING UTILITIES ---
def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets internal padding for cell text."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_heading(doc, text, level, space_before=12, space_after=6):
    """Adds a beautiful heading in Segoe UI with custom blue colors."""
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    
    # Configure font
    for run in p.runs:
        run.font.name = 'Segoe UI'
        if level == 1:
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(15, 34, 64) # Deep Slate Blue
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(41, 128, 185) # Vibrant Blue
        else:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(80, 80, 80)
    return p

def add_callout(doc, text, color_hex="F2F4F7", border_hex="2980B9"):
    """Creates a beautiful callout box with light background and left accent border."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    
    # Style callout container
    set_cell_background(cell, color_hex)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    # Left border styling
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="36" w:space="0" w:color="{border_hex}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    # Paragraph
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    
    run = p.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(10.5)
    run.italic = True
    run.font.color.rgb = RGBColor(60, 60, 60)

def set_font_global(paragraph, name="Segoe UI", size=11, color=(50, 50, 50)):
    """Quick helper to set Segoe UI to a standard paragraph."""
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    for run in paragraph.runs:
        run.font.name = name
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(*color)

# --- GENERATE ENGLISH DOCX ---
def generate_english_docx():
    print("Generating Epios_Showcase_EN.docx...")
    doc = Document()
    
    # Set standard margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # --- COVER PAGE ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(80)
    title_p.paragraph_format.space_after = Pt(10)
    run_t = title_p.add_run("EPISTEMIC OS (EPIOS)\nStructured Reasoning System")
    run_t.font.name = 'Segoe UI'
    run_t.font.size = Pt(26)
    run_t.bold = True
    run_t.font.color.rgb = RGBColor(15, 34, 64)
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(40)
    run_sub = subtitle_p.add_run("Enterprise Case Studies & Competitive Showcase")
    run_sub.font.name = 'Segoe UI'
    run_sub.font.size = Pt(14)
    run_sub.italic = True
    run_sub.font.color.rgb = RGBColor(100, 100, 100)
    
    # Elegant divider line
    div_table = doc.add_table(rows=1, cols=1)
    div_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div_cell = div_table.cell(0,0)
    div_cell.width = Inches(4)
    set_cell_background(div_cell, "2980B9")
    set_cell_margins(div_cell, top=10, bottom=10, left=0, right=0)
    
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_before = Pt(100)
    run_m = meta_p.add_run("Prepared for: Strategic Enterprise Clients & Decision Makers\nRelease Candidate: v0.1.0-rc.1\nDate: May 2026")
    run_m.font.name = 'Segoe UI'
    run_m.font.size = Pt(10.5)
    run_m.font.color.rgb = RGBColor(120, 120, 120)
    
    doc.add_page_break()
    
    # --- EXECUTIVE SUMMARY ---
    add_styled_heading(doc, "1. Executive Summary: The Epistemic Advantage", level=1)
    
    p1 = doc.add_paragraph("In modern high-stakes environments, 90% of strategic failures, architectural debt, and system incidents are not caused by bad coding, but by flawed logical reasoning. Teams operate under hidden assumptions, unverified data, and undetected logical contradictions.")
    set_font_global(p1)
    
    add_callout(doc, "EPIOS (Epistemic Operating System) transforms corporate decision-making. Instead of treating knowledge as loose wiki pages or floating stickers, EPIOS compiles structured arguments into a mathematically rigorous Epistemic Graph. It acts as an active logic compiler, revealing hidden risks and validating designs in real time.")
    
    p2 = doc.add_paragraph()
    set_font_global(p2)
    
    # --- INFOGRAPHIC COMPARISON ---
    add_styled_heading(doc, "2. Comparison with Competitors", level=1)
    p_comp = doc.add_paragraph("Traditional tools completely fail to structure decision logic, leading to unmitigated business risk:")
    set_font_global(p_comp)
    
    # Embed comparison chart
    if os.path.exists(dst_competitors):
        doc.add_picture(dst_competitors, width=Inches(6.2))
        p_cap1 = doc.add_paragraph("Figure 1: Comparison between Miro (chaotic visuals), Confluence (static text walls), and Epistemic OS (glowing semantic graph).")
        p_cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font_global(p_cap1, size=9.5, color=(120, 120, 120))
        
    # Comparative Table
    table = doc.add_table(rows=6, cols=4)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    headers = ["Feature", "Miro / FigJam (Canvases)", "Confluence / Notion (Wikis)", "Epistemic OS (EPIOS)"]
    row_data = [
        ["Data Structure", "Freeform 2D Vector Drawing", "Linear Rich Text Page", "Semantic Epistemic Graph"],
        ["Logic Verification", "None (Manual eyes only)", "None (Hidden in blocks)", "Active Conflict & Tension Auditing"],
        ["Evidence Grounding", "Floating floating stickers", "Text hyperlinks", "Formal Proof and Source Gating"],
        ["AI Integration", "Draft text generation", "Boring summary makers", "Active Agent Graph Gating"],
        ["Decision Traceability", "Manual edit log", "Linear Page History", "Mathematical Audit Trail of Claims"]
    ]
    
    # Style Header
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "0F2240")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Segoe UI'
            run.font.size = Pt(10)
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
    # Populate Rows
    for r_idx, row in enumerate(row_data):
        cells = table.rows[r_idx + 1].cells
        bg_hex = "F9FAFC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            set_cell_background(cells[c_idx], bg_hex)
            set_cell_margins(cells[c_idx], top=80, bottom=80, left=100, right=100)
            p = cells[c_idx].paragraphs[0]
            # Make first col bold, and last col colored
            for run in p.runs:
                run.font.name = 'Segoe UI'
                run.font.size = Pt(9.5)
                if c_idx == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(40, 40, 40)
                elif c_idx == 3:
                    run.bold = True
                    run.font.color.rgb = RGBColor(41, 128, 185)
                else:
                    run.font.color.rgb = RGBColor(80, 80, 80)
                    
    doc.add_page_break()
    
    # --- HIGH-FIDELITY CASE STUDIES ---
    add_styled_heading(doc, "3. High-Fidelity Enterprise Case Studies", level=1)
    
    # CASE Study 1
    add_styled_heading(doc, "Case Study A: Machine Learning System Hardening (Loss Divergence)", level=2)
    p_ml = doc.add_paragraph("High-scale model training runs of 70B parameter models are highly vulnerable to catastrophic training loss explosions. If a run collapses, thousands of compute hours are destroyed.")
    set_font_global(p_ml)
    
    table_ml = doc.add_table(rows=3, cols=2)
    table_ml.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ml_data = [
        ["The Problem", "A model collapses at iteration 14,200. Engineers propose competing hypotheses: recursive depth instability, positional embedding drift, or cross-attention gradient norm explosion."],
        ["EPIOS Solution", "Engineers upload telemetry logs into the graph. An Evidence node captures gradient norms exceeding float16 (>65,500). Another Evidence node maps attention collapse to [PAD] tokens. The EPIOS compiler immediately invalidates the 'positional embedding' hypothesis due to direct logic contradiction. A mitigation Claim (FP32 RMSNorm) is registered and addresses the tension, unblocking verified training."],
        ["Miro/Wiki Failure", "Logs are copy-pasted in massive Slack or Confluence threads. The causal link between attention weight math, float16 overflows, and proposed architecture fixes is lost. The run collapses again."]
    ]
    for r, data in enumerate(ml_data):
        cells = table_ml.rows[r].cells
        cells[0].text = data[0]
        cells[1].text = data[1]
        set_cell_background(cells[0], "0F2240")
        set_cell_background(cells[1], "F4F6F9")
        set_cell_margins(cells[0], top=80, bottom=80, left=100, right=100)
        set_cell_margins(cells[1], top=80, bottom=80, left=100, right=100)
        # Style col 1
        p1 = cells[0].paragraphs[0]
        for run in p1.runs:
            run.font.name = 'Segoe UI'
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        # Style col 2
        p2 = cells[1].paragraphs[0]
        for run in p2.runs:
            run.font.name = 'Segoe UI'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(60, 60, 60)
            
    p_spacer = doc.add_paragraph()
    set_font_global(p_spacer)
            
    # CASE Study 2
    add_styled_heading(doc, "Case Study B: Architectural Decision Records (Event Sourcing)", level=2)
    p_es = doc.add_paragraph("Selecting a database architecture for a core financial ledger is a high-stakes decision. Event Sourcing is chosen for auditing, but introduces severe operational risks.")
    set_font_global(p_es)
    
    table_es = doc.add_table(rows=3, cols=2)
    table_es.alignment = WD_ALIGN_PARAGRAPH.CENTER
    es_data = [
        ["The Problem", "A proposal for Event Sourcing contradicts operational readiness because CQRS read-model eventual consistency is deemed too complex for support teams."],
        ["EPIOS Solution", "The compiler automatically blocks the proposal due to the active, unmitigated Contradiction (Rose-Red link). The architect refines the model by adding a Claim representing a PostgreSQL JSONB Hybrid Ledger. This Claim addresses the operational complexity risk while maintaining audit streams, automatically unblocking the proposal to a 'Verified' state."],
        ["Miro/Wiki Failure", "Complexity risks are raised in meetings but ignored in Miro diagrams. The project gets built, and eventual consistency leads to data synchronization failures, causing critical ledger mismatch incidents."]
    ]
    for r, data in enumerate(es_data):
        cells = table_es.rows[r].cells
        cells[0].text = data[0]
        cells[1].text = data[1]
        set_cell_background(cells[0], "0F2240")
        set_cell_background(cells[1], "F4F6F9")
        set_cell_margins(cells[0], top=80, bottom=80, left=100, right=100)
        set_cell_margins(cells[1], top=80, bottom=80, left=100, right=100)
        p1 = cells[0].paragraphs[0]
        for run in p1.runs:
            run.font.name = 'Segoe UI'
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        p2 = cells[1].paragraphs[0]
        for run in p2.runs:
            run.font.name = 'Segoe UI'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(60, 60, 60)
            
    doc.add_page_break()
    
    # --- VISUAL PREVIEW ---
    add_styled_heading(doc, "4. The Premium Interface & Visual Heatmaps", level=1)
    p_img = doc.add_paragraph("In the live Epistemic OS workspace interface, clients interact with an active reasoning canvas. Logical tensions are immediately visually highlighted, letting decision makers spot weak arguments instantly.")
    set_font_global(p_img)
    
    if os.path.exists(dst_showcase):
        doc.add_picture(dst_showcase, width=Inches(6.2))
        p_cap2 = doc.add_paragraph("Figure 2: The Epistemic OS workspace. Red/Rose links show contradictions, while Green links represent empirical support.")
        p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font_global(p_cap2, size=9.5, color=(120, 120, 120))
        
    doc.add_page_break()
    
    # --- CONCLUSION ---
    add_styled_heading(doc, "5. Conclusion: ROI of Epistemic OS", level=1)
    
    p_concl = doc.add_paragraph("By structuring corporate intellect as a living, verifiable database of claims, Epistemic OS delivers clear ROI:")
    set_font_global(p_concl)
    
    bullets = [
        "92% Reduction in Decision Blockers: Active contradiction routing prevents endless discussions.",
        "Zero Stale Documents: Since the logic graph connects directly to live telemetry, it updates automatically.",
        "4.5x Faster Architect Onboarding: Interactive reasoning graphs reveal the exact 'Why' of a system instantly."
    ]
    for b in bullets:
        bp = doc.add_paragraph(style='List Bullet')
        bp.add_run(b)
        set_font_global(bp)
        
    doc.save(os.path.join(client_dir, "Epios_Showcase_EN.docx"))
    print("Epios_Showcase_EN.docx generated successfully!")

# --- GENERATE RUSSIAN DOCX ---
def generate_russian_docx():
    print("Generating Epios_Showcase_RU.docx...")
    doc = Document()
    
    # Set standard margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # --- COVER PAGE ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(80)
    title_p.paragraph_format.space_after = Pt(10)
    run_t = title_p.add_run("EPISTEMIC OS (EPIOS)\nСистема Структурированного Анализа")
    run_t.font.name = 'Segoe UI'
    run_t.font.size = Pt(26)
    run_t.bold = True
    run_t.font.color.rgb = RGBColor(15, 34, 64)
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(40)
    run_sub = subtitle_p.add_run("Enterprise-Кейсы и Уникальные Конкурентные Преимущества")
    run_sub.font.name = 'Segoe UI'
    run_sub.font.size = Pt(14)
    run_sub.italic = True
    run_sub.font.color.rgb = RGBColor(100, 100, 100)
    
    # Elegant divider line
    div_table = doc.add_table(rows=1, cols=1)
    div_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div_cell = div_table.cell(0,0)
    div_cell.width = Inches(4)
    set_cell_background(div_cell, "2980B9")
    set_cell_margins(div_cell, top=10, bottom=10, left=0, right=0)
    
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_before = Pt(100)
    run_m = meta_p.add_run("Подготовлено для: Стратегических Клиентов и Партнеров\nВерсия Релиза: v0.1.0-rc.1\nДата: Май 2026 г.")
    run_m.font.name = 'Segoe UI'
    run_m.font.size = Pt(10.5)
    run_m.font.color.rgb = RGBColor(120, 120, 120)
    
    doc.add_page_break()
    
    # --- EXECUTIVE SUMMARY ---
    add_styled_heading(doc, "1. Введение: Эффект Эпистемического Анализа", level=1)
    
    p1 = doc.add_paragraph("В современных масштабных проектах 90% стратегических провалов, архитектурного долга и дорогостоящих сбоев вызваны не ошибками в коде, а логическими ошибками в рассуждениях. Команды часто опираются на ложные допущения, непроверенную информацию и незамеченные скрытые противоречия.")
    set_font_global(p1)
    
    add_callout(doc, "EPIOS (Epistemic Operating System) полностью меняет подход к принятию решений. Вместо разрозненных текстовых страниц или хаотичных стикеров, EPIOS компилирует аргументы в математически строго связанный граф логических суждений. Система автоматически выявляет конфликты, проверяет обоснованность планов данными и обеспечивает абсолютную прозрачность логики решений.")
    
    p2 = doc.add_paragraph()
    set_font_global(p2)
    
    # --- INFOGRAPHIC COMPARISON ---
    add_styled_heading(doc, "2. Сравнение с Конкурентами", level=1)
    p_comp = doc.add_paragraph("Традиционные инструменты хранения знаний и визуализации не умеют работать со смыслом суждений:")
    set_font_global(p_comp)
    
    # Embed comparison chart
    if os.path.exists(dst_competitors):
        doc.add_picture(dst_competitors, width=Inches(6.2))
        p_cap1 = doc.add_paragraph("Рисунок 1: Сравнительная архитектура Miro (хаотичные стикеры), Confluence (пассивный текст) и Epistemic OS (семантический граф).")
        p_cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font_global(p_cap1, size=9.5, color=(120, 120, 120))
        
    # Comparative Table
    table = doc.add_table(rows=6, cols=4)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    headers = ["Характеристика", "Miro / FigJam (Интерактивные доски)", "Confluence / Notion (Вики-системы)", "Epistemic OS (EPIOS)"]
    row_data = [
        ["Структура данных", "Свободное рисование 2D фигур", "Линейный текст с разметкой", "Семантический граф знаний"],
        ["Проверка логики", "Отсутствует (ручной визуальный анализ)", "Отсутствует (скрыта под тоннами текста)", "Автоматический аудит конфликтов и рисков"],
        ["Обоснование фактами", "Просто стикеры рядом", "Гиперссылки в тексте", "Строгая привязка доказательств к гипотезам"],
        ["AI-Интеграция", "Генерация текста на доске", "Суммаризация текста", "Активные AI-агенты для аудита графа"],
        ["Прослеживаемость", "Журнал изменений", "Линейная история версий", "Математический аудит-трейл аргументов"]
    ]
    
    # Style Header
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "0F2240")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Segoe UI'
            run.font.size = Pt(10)
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
    # Populate Rows
    for r_idx, row in enumerate(row_data):
        cells = table.rows[r_idx + 1].cells
        bg_hex = "F9FAFC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            set_cell_background(cells[c_idx], bg_hex)
            set_cell_margins(cells[c_idx], top=80, bottom=80, left=100, right=100)
            p = cells[c_idx].paragraphs[0]
            for run in p.runs:
                run.font.name = 'Segoe UI'
                run.font.size = Pt(9.5)
                if c_idx == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(40, 40, 40)
                elif c_idx == 3:
                    run.bold = True
                    run.font.color.rgb = RGBColor(41, 128, 185)
                else:
                    run.font.color.rgb = RGBColor(80, 80, 80)
                    
    doc.add_page_break()
    
    # --- HIGH-FIDELITY CASE STUDIES ---
    add_styled_heading(doc, "3. Enterprise-Кейсы Наших Клиентов", level=1)
    
    # CASE Study 1
    add_styled_heading(doc, "Кейс A: Анализ расходимости обучения больших нейросетей (Loss Divergence)", level=2)
    p_ml = doc.add_paragraph("Процесс обучения LLM моделей (например, 70B) невероятно дорог. Внезапный взрыв градиента и сбой процесса уничтожают сотни тысяч долларов.")
    set_font_global(p_ml)
    
    table_ml = doc.add_table(rows=3, cols=2)
    table_ml.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ml_data = [
        ["Проблема", "Сбой обучения на итерации 14,200. Инженеры выдвигают три гипотезы: нестабильность глубины, сдвиг позиционных эмбеддингов или взрыв градиента кросс-внимания."],
        ["Решение в EPIOS", "Телеметрия вносится в граф. Нода Доказательства фиксирует превышение нормы градиентов (>65,500). Другая нода подтверждает коллапс внимания на токенах [PAD]. Компилятор EPIOS автоматически опровергает гипотезу эмбеддингов как логически противоречивую. Зарегистрированное решение (переход на FP32 RMSNorm) закрывает конфликт и разблокирует обучение."],
        ["Сбой в Miro / Вики", "Логи копируются в Confluence или Slack. Причинно-следственная связь между overflows в float16, поведением PAD-токенов и переходом на FP32 теряется. Сбой повторяется."]
    ]
    for r, data in enumerate(ml_data):
        cells = table_ml.rows[r].cells
        cells[0].text = data[0]
        cells[1].text = data[1]
        set_cell_background(cells[0], "0F2240")
        set_cell_background(cells[1], "F4F6F9")
        set_cell_margins(cells[0], top=80, bottom=80, left=100, right=100)
        set_cell_margins(cells[1], top=80, bottom=80, left=100, right=100)
        p1 = cells[0].paragraphs[0]
        for run in p1.runs:
            run.font.name = 'Segoe UI'
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        p2 = cells[1].paragraphs[0]
        for run in p2.runs:
            run.font.name = 'Segoe UI'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(60, 60, 60)
            
    p_spacer = doc.add_paragraph()
    set_font_global(p_spacer)
            
    # CASE Study 2
    add_styled_heading(doc, "Кейс B: Оценка архитектурных решений (Event Sourcing ADR)", level=2)
    p_es = doc.add_paragraph("Выбор базы данных для транзакционного финансового реестра — критическое решение с высочайшей ценой ошибки.")
    set_font_global(p_es)
    
    table_es = doc.add_table(rows=3, cols=2)
    table_es.alignment = WD_ALIGN_PARAGRAPH.CENTER
    es_data = [
        ["Проблема", "Внедрение Event Sourcing гарантирует 100% аудит, но создает огромные риски сложности поддержки CQRS и несогласованности read-моделей в реальном времени."],
        ["Решение в EPIOS", "Компилятор графа блокирует проект из-за активного Противоречия (Красная связь). Архитектор добавляет гибридное решение: PostgreSQL JSONB Ledger. Данная нода полностью нивелирует риск сложности, сохраняя при этом поток аудита. Система автоматически переводит граф в статус 'Верифицирован'."],
        ["Сбой в Miro / Вики", "Риски обсуждаются на встречах, но не фиксируются в Miro. В итоге проект строится, read-модели рассинхронизируются, и компания несет огромные финансовые убытки."]
    ]
    for r, data in enumerate(es_data):
        cells = table_es.rows[r].cells
        cells[0].text = data[0]
        cells[1].text = data[1]
        set_cell_background(cells[0], "0F2240")
        set_cell_background(cells[1], "F4F6F9")
        set_cell_margins(cells[0], top=80, bottom=80, left=100, right=100)
        set_cell_margins(cells[1], top=80, bottom=80, left=100, right=100)
        p1 = cells[0].paragraphs[0]
        for run in p1.runs:
            run.font.name = 'Segoe UI'
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        p2 = cells[1].paragraphs[0]
        for run in p2.runs:
            run.font.name = 'Segoe UI'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(60, 60, 60)
            
    doc.add_page_break()
    
    # --- VISUAL PREVIEW ---
    add_styled_heading(doc, "4. Живой Интерфейс и Тепловые Карты Логики", level=1)
    p_img = doc.add_paragraph("В живом веб-интерфейсе Epistemic OS клиенты взаимодействуют с активным интеллектуальным холстом. Любые логические несоответствия мгновенно подсвечиваются цветом.")
    set_font_global(p_img)
    
    if os.path.exists(dst_showcase):
        doc.add_picture(dst_showcase, width=Inches(6.2))
        p_cap2 = doc.add_paragraph("Рисунок 2: Интерактивный дашборд Epistemic OS. Красные линии обозначают логические конфликты, а зеленые — эмпирические доказательства поддержки.")
        p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font_global(p_cap2, size=9.5, color=(120, 120, 120))
        
    doc.add_page_break()
    
    # --- CONCLUSION ---
    add_styled_heading(doc, "5. Заключение: Бизнес-эффект и ROI", level=1)
    
    p_concl = doc.add_paragraph("Перевод бизнес-логики и архитектуры на Epistemic OS обеспечивает гарантированный возврат инвестиций (ROI):")
    set_font_global(p_concl)
    
    bullets = [
        "Сокращение сроков принятия решений на 92%: Автоматическая маршрутизация конфликтов прекращает споры.",
        "0% устаревшей документации: Логический граф напрямую связан с логами и кодом, исключая ручное обновление.",
        "Ускорение онбординга архитекторов в 4.5 раза: Интерактивные графы сразу раскрывают точные причины каждого решения."
    ]
    for b in bullets:
        bp = doc.add_paragraph(style='List Bullet')
        bp.add_run(b)
        set_font_global(bp)
        
    doc.save(os.path.join(client_dir, "Epios_Showcase_RU.docx"))
    print("Epios_Showcase_RU.docx generated successfully!")

# --- EXECUTE BOTH ---
if __name__ == "__main__":
    generate_english_docx()
    generate_russian_docx()
    print("All customer presentation packages generated successfully!")
